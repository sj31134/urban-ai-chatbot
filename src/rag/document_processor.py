"""
도시정비 법령 문서 처리 파이프라인
PDF, DOCX 등 법령 문서를 구조화하여 Neo4j 그래프로 변환
"""

# 필요한 라이브러리 import
import os  # 운영체제 관련 기능
import re  # 정규표현식
import json  # JSON 데이터 처리
import logging  # 로그 기록
from typing import Dict, List, Optional, Tuple, Any  # 타입 힌트
from datetime import datetime  # 날짜/시간 처리
from pathlib import Path  # 경로 처리

import PyPDF2  # PDF 처리 라이브러리 (구버전)
import pypdf  # PDF 처리 라이브러리 (신버전)
from docx import Document  # DOCX 문서 처리
from bs4 import BeautifulSoup  # HTML 파싱
from langchain.text_splitter import RecursiveCharacterTextSplitter  # 텍스트 분할
from langchain.schema import Document as LangChainDocument  # LangChain 문서 클래스

from src.graph.legal_graph import LegalGraphManager  # 그래프 관리자

# 로깅 설정
logging.basicConfig(level=logging.INFO)  # INFO 레벨 로깅 설정
logger = logging.getLogger(__name__)  # 현재 모듈용 로거 생성


class LegalDocumentProcessor:
    """법령 문서 처리 및 그래프 변환 클래스"""
    
    def __init__(self, graph_manager: LegalGraphManager, schema_path: str = "config/legal_schema.json"):
        """
        문서 처리기 초기화
        
        Args:
            graph_manager: Neo4j 그래프 관리자
            schema_path: 법령 스키마 설정 파일 경로
        """
        self.graph_manager = graph_manager  # 그래프 관리자 저장
        self.schema = self._load_schema(schema_path)  # 스키마 로드
        self.text_splitter = self._create_text_splitter()  # 텍스트 분할기 생성
        
        # 법령 구조 패턴
        self.article_pattern = re.compile(r'제(\d+)조(의\d+)?')  # 조문 패턴 (제N조, 제N조의N)
        self.section_pattern = re.compile(r'제(\d+)편|제(\d+)장|제(\d+)절')  # 편장절 패턴
        self.subsection_pattern = re.compile(r'①|②|③|④|⑤|⑥|⑦|⑧|⑨|⑩')  # 항 패턴
        self.item_pattern = re.compile(r'\d+\.|가\.|나\.|다\.|라\.|마\.|바\.|사\.|아\.|자\.|차\.|카\.|타\.|파\.|하\.')  # 호 패턴
        
    def _load_schema(self, schema_path: str) -> Dict:
        """법령 스키마 로드"""
        try:
            with open(schema_path, 'r', encoding='utf-8') as file:  # 스키마 파일 열기
                return json.load(file)  # JSON 파일 로드
        except FileNotFoundError:  # 파일을 찾을 수 없으면
            logger.warning(f"스키마 파일을 찾을 수 없습니다: {schema_path}")  # 경고 로그
            return {}  # 빈 딕셔너리 반환
    
    def _create_text_splitter(self) -> RecursiveCharacterTextSplitter:
        """법령 구조를 고려한 텍스트 분할기 생성"""
        chunk_config = self.schema.get("document_processing", {}).get("chunk_strategy", {})  # 청크 설정 가져오기
        
        return RecursiveCharacterTextSplitter(  # 재귀적 텍스트 분할기 생성
            chunk_size=chunk_config.get("chunk_size", 512),  # 청크 크기
            chunk_overlap=chunk_config.get("chunk_overlap", 50),  # 청크 겹침
            separators=chunk_config.get("separators", ["제", "조", "항", "호", "목"]),  # 분할 기준
            keep_separator=True,  # 분할 기준 문자 유지
            add_start_index=True  # 시작 인덱스 추가
        )
    
    def process_document(self, file_path: str, law_code: str = None) -> Dict:
        """문서 타입에 따라 자동으로 처리"""
        file_extension = Path(file_path).suffix.lower()  # 파일 확장자 추출
        
        if law_code is None:  # 법령 코드가 없으면
            law_code = self._infer_law_code(file_path)  # 파일명에서 추론
        
        if file_extension == '.pdf':  # PDF 파일인 경우
            return self.process_pdf_document(file_path, law_code)
        elif file_extension == '.docx':  # DOCX 파일인 경우
            return self.process_docx_document(file_path, law_code)
        elif file_extension == '.doc':  # DOC 파일인 경우
            return self.process_doc_document(file_path, law_code)
        elif file_extension == '.hwp':  # HWP 파일인 경우
            return self.process_hwp_document(file_path, law_code)
        elif file_extension == '.txt':  # TXT 파일인 경우
            return self.process_txt_document(file_path, law_code)
        else:  # 지원하지 않는 형식인 경우
            logger.error(f"지원하지 않는 파일 형식: {file_extension}")
            return {"success": False, "error": f"Unsupported file format: {file_extension}"}
    
    def process_pdf_document(self, file_path: str, law_code: str) -> Dict:
        """PDF 법령 문서 처리"""
        logger.info(f"PDF 문서 처리 시작: {file_path}")  # 처리 시작 로그
        
        # PDF 텍스트 추출
        text_content = self._extract_pdf_text(file_path)  # PDF에서 텍스트 추출
        
        # 법령 정보 추출
        law_info = self._extract_law_metadata(text_content, law_code)  # 법령 메타데이터 추출
        
        # 조문별 분할
        articles = self._extract_articles(text_content, law_code)  # 조문별로 분할
        
        # Neo4j에 저장
        result = self._save_to_graph(law_info, articles)  # 그래프 데이터베이스에 저장
        
        logger.info(f"PDF 처리 완료: {len(articles)}개 조문 처리")  # 처리 완료 로그
        return result  # 처리 결과 반환
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """PDF에서 텍스트 추출"""
        try:
            with open(file_path, 'rb') as file:  # PDF 파일을 바이너리 모드로 열기
                # pypdf 사용 (최신 버전)
                reader = pypdf.PdfReader(file)  # PDF 리더 생성
                text_content = ""  # 텍스트 내용 초기화
                
                for page in reader.pages:  # 각 페이지에 대해
                    text_content += page.extract_text() + "\n"  # 텍스트 추출하여 누적
                
                return text_content.strip()  # 앞뒤 공백 제거하여 반환
                
        except Exception as e:  # 예외 발생 시
            logger.error(f"PDF 텍스트 추출 실패: {e}")  # 에러 로그
            # 대안으로 PyPDF2 시도
            try:
                with open(file_path, 'rb') as file:  # PDF 파일을 바이너리 모드로 열기
                    reader = PyPDF2.PdfReader(file)  # PyPDF2 리더 생성
                    text_content = ""  # 텍스트 내용 초기화
                    
                    for page in reader.pages:  # 각 페이지에 대해
                        text_content += page.extract_text() + "\n"  # 텍스트 추출하여 누적
                    
                    return text_content.strip()  # 앞뒤 공백 제거하여 반환
            except Exception as e2:  # 두 번째 시도도 실패 시
                logger.error(f"PDF 텍스트 추출 실패 (PyPDF2): {e2}")  # 에러 로그
                return ""  # 빈 문자열 반환
    
    def _extract_law_metadata(self, text_content: str, law_code: str) -> Dict:
        """법령 메타데이터 추출"""
        law_info = {  # 법령 정보 딕셔너리 초기화
            "law_id": law_code,  # 법령 ID
            "category": "법률",  # 법령 분류
            "status": "시행",  # 시행 상태
            "effective_date": "2024-01-01",  # 기본 시행일
            "last_updated": datetime.now().isoformat(),  # 최종 업데이트 시간
            "created_at": datetime.now().isoformat()  # 생성 시간
        }
        
        # 법령명 추출
        title_patterns = [  # 법령명 패턴 목록
            r'([가-힣\s]+법)\s*\n',  # ~법
            r'([가-힣\s]+에\s*관한\s*법률)',  # ~에 관한 법률
            r'([가-힣\s]+특례법)'  # ~특례법
        ]
        
        for pattern in title_patterns:  # 각 패턴에 대해
            match = re.search(pattern, text_content[:500])  # 첫 500자에서 패턴 검색
            if match:  # 매칭되면
                law_info["name"] = match.group(1).strip()  # 법령명 추출
                break  # 첫 번째 매칭에서 중단
        
        # 기본 법령명 설정 (패턴 매칭 실패 시)
        if "name" not in law_info:
            law_info["name"] = f"법령_{law_code}"
        
        # 시행일 추출
        date_patterns = [  # 시행일 패턴 목록
            r'시행일\s*[:：]\s*(\d{4})\s*년\s*(\d{1,2})\s*월\s*(\d{1,2})\s*일',  # 시행일: YYYY년 MM월 DD일
            r'(\d{4})\.\s*(\d{1,2})\.\s*(\d{1,2})\.\s*시행',  # YYYY.MM.DD. 시행
            r'(\d{4})-(\d{2})-(\d{2})\s*시행'  # YYYY-MM-DD 시행
        ]
        
        for pattern in date_patterns:  # 각 패턴에 대해
            match = re.search(pattern, text_content[:1000])  # 첫 1000자에서 패턴 검색
            if match:  # 매칭되면
                year, month, day = match.groups()  # 연월일 추출
                law_info["effective_date"] = f"{year}-{month.zfill(2)}-{day.zfill(2)}"  # 표준 날짜 형식으로 변환
                break  # 첫 번째 매칭에서 중단
        
        return law_info  # 법령 정보 반환
    
    def _extract_articles(self, text_content: str, law_code: str) -> List[Dict]:
        """조문별로 텍스트 분할 및 메타데이터 추출"""
        articles = []  # 조문 리스트 초기화
        
        # 조문 분할
        article_blocks = self._split_by_articles(text_content)  # 조문별로 분할
        
        current_section = ""  # 현재 절
        current_chapter = ""  # 현재 장
        
        for block in article_blocks:  # 각 조문 블록에 대해
            article_info = self._parse_article_block(block, law_code)  # 조문 정보 파싱
            
            if article_info:  # 조문 정보가 있으면
                # 편장절 정보 업데이트
                section_match = self.section_pattern.search(block[:100])  # 블록 앞부분에서 편장절 패턴 검색
                if section_match:  # 매칭되면
                    current_section = section_match.group(0)  # 현재 절 업데이트
                
                article_info["section"] = current_section  # 절 정보 추가
                article_info["chapter"] = current_chapter  # 장 정보 추가
                articles.append(article_info)  # 조문 리스트에 추가
        
        return articles  # 조문 리스트 반환
    
    def _split_by_articles(self, text_content: str) -> List[str]:
        """조문 기준으로 텍스트 분할"""
        # 조문 시작 패턴으로 분할
        article_starts = []  # 조문 시작 위치 리스트
        for match in self.article_pattern.finditer(text_content):  # 조문 패턴으로 검색
            article_starts.append(match.start())  # 시작 위치 추가
        
        if not article_starts:  # 조문이 없으면
            return [text_content]  # 전체 텍스트 반환
        
        # 각 조문 블록 추출
        blocks = []  # 조문 블록 리스트
        for i, start in enumerate(article_starts):  # 각 조문 시작 위치에 대해
            if i < len(article_starts) - 1:  # 마지막이 아니면
                end = article_starts[i + 1]  # 다음 조문 시작 위치를 끝으로
                blocks.append(text_content[start:end])  # 해당 범위 추출
            else:  # 마지막이면
                blocks.append(text_content[start:])  # 끝까지 추출
        
        return blocks  # 조문 블록들 반환
    
    def _parse_article_block(self, block: str, law_code: str) -> Optional[Dict]:
        """조문 블록 파싱"""
        # 조문 번호 추출
        article_match = self.article_pattern.search(block[:50])  # 블록 앞부분에서 조문 번호 검색
        if not article_match:  # 조문 번호가 없으면
            return None  # None 반환
        
        article_number = article_match.group(0)  # 조문 번호 추출
        
        # 조문 내용 추출 (조문 번호 이후부터)
        content_start = article_match.end()  # 조문 번호 이후 시작 위치
        content = block[content_start:].strip()  # 조문 내용 추출
        
        # 조문 내용에서 불필요한 부분 제거
        content = re.sub(r'\s+', ' ', content)  # 연속된 공백을 하나로
        content = content.replace('\n', ' ').strip()  # 줄바꿈을 공백으로 변경
        
        # 조문 내용이 너무 긴 경우 적절히 분할 또는 요약
        if len(content) > 20000:  # 20KB 제한
            logger.warning(f"조문 {article_number} 내용이 너무 김 ({len(content)} 문자), 분할 처리")
            # 첫 10,000자만 사용하고 나머지는 요약 표시
            content = content[:10000] + "... [내용이 길어 일부 생략됨]"
        
        # 하위 조항 추출
        subsections = self._extract_subsections(content)  # 항 추출
        items = self._extract_items(content)  # 호 추출
        
        # 참조 조문 찾기
        references = self._find_article_references(content)  # 참조 조문 검색
        
        return {  # 조문 정보 딕셔너리 반환
            "law_id": law_code,  # 법령 ID
            "article_number": article_number,  # 조문 번호
            "content": content,  # 조문 내용
            "section": "",  # 절 정보 (나중에 업데이트)
            "subsection": "",  # 하위 절 정보
            "subsections": subsections,  # 하위 항들
            "items": items,  # 하위 호들
            "references": references,  # 참조 조문들
            "last_updated": datetime.now().isoformat()  # 최종 업데이트 시간
        }
    
    def _extract_subsections(self, content: str) -> List[str]:
        """항(①, ②, ...) 추출"""
        subsections = []  # 항 리스트
        parts = self.subsection_pattern.split(content)  # 항 패턴으로 분할
        
        for i in range(1, len(parts)):  # 첫 번째를 제외한 각 부분에 대해
            if i < len(parts):  # 범위 내에 있으면
                # 항 번호와 내용 결합
                subsection_marker = self.subsection_pattern.findall(content)[i-1] if i-1 < len(self.subsection_pattern.findall(content)) else f"⑪{i}"  # 항 번호
                subsection_content = parts[i].strip()  # 항 내용
                
                if subsection_content:  # 내용이 있으면
                    subsections.append(f"{subsection_marker} {subsection_content}")  # 항 리스트에 추가
        
        return subsections  # 항 리스트 반환
    
    def _extract_items(self, content: str) -> List[str]:
        """호(1., 2., 가., 나., ...) 추출"""
        items = []  # 호 리스트
        
        # 숫자 호 패턴
        number_items = re.findall(r'(\d+\.)\s*([^가-하\d\.]+?)(?=\d+\.|$)', content)  # 숫자 호 추출
        for number, item_content in number_items:  # 각 숫자 호에 대해
            items.append(f"{number} {item_content.strip()}")  # 호 리스트에 추가
        
        # 문자 호 패턴
        char_items = re.findall(r'([가-하]\.)\s*([^가-하\d\.]+?)(?=[가-하]\.|$)', content)  # 문자 호 추출
        for char, item_content in char_items:  # 각 문자 호에 대해
            items.append(f"{char} {item_content.strip()}")  # 호 리스트에 추가
        
        return items  # 호 리스트 반환
    
    def _find_article_references(self, content: str) -> List[str]:
        """조문 간 참조 관계 찾기"""
        references = []  # 참조 리스트
        
        # 참조 패턴들
        reference_patterns = [  # 참조 패턴 목록
            r'제(\d+)조(?:의(\d+))?',  # 제N조, 제N조의N
            r'같은\s*조',  # 같은 조
            r'이\s*법',  # 이 법
            r'앞\s*조',  # 앞 조
            r'다음\s*조'  # 다음 조
        ]
        
        for pattern in reference_patterns:  # 각 패턴에 대해
            matches = re.findall(pattern, content)  # 패턴 매칭
            for match in matches:  # 각 매칭에 대해
                if isinstance(match, tuple):  # 튜플이면
                    ref = f"제{match[0]}조"  # 조문 번호 구성
                    if match[1]:  # 하위 조문이 있으면
                        ref += f"의{match[1]}"  # 하위 조문 추가
                    references.append(ref)  # 참조 리스트에 추가
                else:  # 문자열이면
                    references.append(match)  # 그대로 추가
        
        return list(set(references))  # 중복 제거하여 반환
    
    def _save_to_graph(self, law_info: Dict, articles: List[Dict]) -> Dict:
        """그래프 데이터베이스에 법령 정보 저장"""
        try:
            # 법령 노드 생성
            law_id = self.graph_manager.create_law_node(law_info)  # 법령 노드 생성
            logger.info(f"법령 노드 생성 완료: {law_id}")  # 생성 완료 로그
            
            # 조문 노드들 생성
            article_nodes = []  # 조문 노드 리스트
            for article in articles:  # 각 조문에 대해
                try:
                    article_id = self.graph_manager.create_article_node(article)  # 조문 노드 생성
                    
                    # 법령-조문 관계 생성
                    self.graph_manager.create_relationship(  # 관계 생성
                        article_id, law_id, "BELONGS_TO",  # 조문이 법령에 속함
                        {"section": article.get("section", ""), "chapter": article.get("chapter", "")}  # 관계 속성
                    )
                    
                    article_nodes.append({  # 조문 노드 정보 추가
                        "node_id": article_id,
                        "article_number": article["article_number"],
                        "references": article.get("references", [])
                    })
                    
                except Exception as e:  # 조문 노드 생성 실패 시
                    logger.error(f"조문 노드 생성 실패: {article['article_number']} - {e}")  # 에러 로그
            
            # 조문 간 참조 관계 생성
            self._create_cross_references(article_nodes)  # 교차 참조 관계 생성
            
            result = {  # 결과 딕셔너리
                "law_id": law_id,  # 법령 ID
                "total_articles": len(articles),  # 총 조문 수
                "processed_articles": len(article_nodes),  # 처리된 조문 수
                "status": "success"  # 처리 상태
            }
            
            logger.info(f"그래프 저장 완료: {result}")  # 저장 완료 로그
            return result  # 결과 반환
            
        except Exception as e:  # 전체 처리 실패 시
            logger.error(f"그래프 저장 실패: {e}")  # 에러 로그
            return {"status": "error", "message": str(e)}  # 에러 결과 반환
    
    def _create_cross_references(self, article_nodes: List[Dict]):
        """조문 간 교차 참조 관계 생성"""
        for article_node in article_nodes:  # 각 조문 노드에 대해
            references = article_node.get("references", [])  # 참조 조문들 가져오기
            
            for ref in references:  # 각 참조에 대해
                # 참조되는 조문 찾기
                for target_node in article_nodes:  # 각 대상 노드에 대해
                    if target_node["article_number"] == ref:  # 조문 번호가 일치하면
                        try:
                            # 참조 관계 생성
                            self.graph_manager.create_relationship(  # 관계 생성
                                article_node["node_id"],  # 참조하는 조문
                                target_node["node_id"],  # 참조되는 조문
                                "REFERENCES",  # 참조 관계
                                {  # 관계 속성
                                    "reference_type": "direct",  # 직접 참조
                                    "context": "article_reference"  # 조문 참조
                                }
                            )
                            logger.debug(f"참조 관계 생성: {article_node['article_number']} -> {ref}")  # 디버그 로그
                        except Exception as e:  # 관계 생성 실패 시
                            logger.warning(f"참조 관계 생성 실패: {e}")  # 경고 로그
                        break  # 매칭되면 중단
    
    def process_documents_batch(self, data_directory: str) -> Dict:
        """디렉토리 내 모든 문서 일괄 처리"""
        results = {  # 결과 딕셔너리
            "total_files": 0,  # 총 파일 수
            "processed_files": 0,  # 처리된 파일 수
            "failed_files": [],  # 실패한 파일들
            "processing_details": []  # 처리 상세 정보
        }
        
        data_path = Path(data_directory)  # 데이터 디렉토리 경로
        
        # 지원하는 파일 확장자
        supported_extensions = ['.pdf', '.docx', '.doc', '.hwp', '.txt']  # 지원 확장자 목록
        
        for file_path in data_path.rglob('*'):  # 모든 하위 파일에 대해
            if file_path.suffix.lower() in supported_extensions:  # 지원하는 확장자면
                results["total_files"] += 1  # 총 파일 수 증가
                
                try:
                    # 파일명에서 법령 코드 추정
                    law_code = self._infer_law_code(file_path.name)  # 법령 코드 추정
                    
                    # 파일 처리
                    process_result = self.process_document(str(file_path), law_code)  # 문서 처리
                    
                    if process_result.get("status") == "success":  # 처리 성공하면
                        results["processed_files"] += 1  # 처리된 파일 수 증가
                        
                    results["processing_details"].append({  # 처리 상세 정보 추가
                        "file_path": str(file_path),
                        "law_code": law_code,
                        "result": process_result
                    })
                    
                except Exception as e:  # 파일 처리 실패 시
                    logger.error(f"파일 처리 실패: {file_path} - {e}")  # 에러 로그
                    results["failed_files"].append(str(file_path))  # 실패 파일에 추가
        
        logger.info(f"일괄 처리 완료: {results['processed_files']}/{results['total_files']} 성공")  # 완료 로그
        return results  # 결과 반환
    
    def _infer_law_code(self, filename: str) -> str:
        """파일명에서 법령 코드 추정"""
        # 파일명 기반 법령 코드 매핑
        code_mappings = {  # 파일명-법령코드 매핑
            "도시정비법": "urban_redevelopment_law",  # 도시정비법
            "소규모주택": "small_housing_law",  # 소규모주택정비법
            "빈집정비": "vacant_house_law",  # 빈집정비특례법
            "안양시": "anyang_ordinance",  # 안양시 조례
            "성남시": "seongnam_ordinance",  # 성남시 조례
            "용인시": "yongin_ordinance",  # 용인시 조례
            "서울": "seoul_ordinance"  # 서울시 조례
        }
        
        filename_lower = filename.lower()  # 파일명을 소문자로 변환
        
        for keyword, code in code_mappings.items():  # 각 매핑에 대해
            if keyword in filename:  # 키워드가 파일명에 포함되어 있으면
                # 중복 방지를 위해 타임스탬프 추가
                import time
                timestamp = str(int(time.time()))[-6:]  # 마지막 6자리
                return f"{code}_{timestamp}"  # 타임스탬프가 포함된 고유 코드 반환
        
        # 기본 코드 생성 (파일명 기반)
        base_name = Path(filename).stem  # 확장자 제거한 파일명
        clean_name = re.sub(r'[^\w]', '_', base_name.lower())  # 특수문자를 언더스코어로 변경
        
        # 중복 방지를 위해 타임스탬프 추가
        import time
        timestamp = str(int(time.time()))[-6:]
        return f"{clean_name}_{timestamp}"
    
    def process_txt_document(self, file_path: str, law_code: str) -> Dict:
        """TXT 법령 문서 처리"""
        logger.info(f"TXT 문서 처리 시작: {file_path}")
        
        # TXT 텍스트 추출
        text_content = self._extract_txt_text(file_path)
        
        if not text_content:  # 텍스트가 추출되지 않았으면
            return {"success": False, "error": "Failed to extract text from TXT file"}
        
        # 법령 정보 추출
        law_info = self._extract_law_metadata(text_content, law_code)
        
        # 조문별 분할
        articles = self._extract_articles(text_content, law_code)
        
        # Neo4j에 저장
        result = self._save_to_graph(law_info, articles)
        
        logger.info(f"TXT 처리 완료: {len(articles)}개 조문 처리")
        return result
    
    def _extract_txt_text(self, file_path: str) -> str:
        """TXT 파일에서 텍스트 추출"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:  # UTF-8로 파일 열기
                return file.read().strip()  # 텍스트 읽고 앞뒤 공백 제거
        except UnicodeDecodeError:  # UTF-8 실패 시
            try:
                with open(file_path, 'r', encoding='cp949') as file:  # CP949로 재시도
                    return file.read().strip()
            except Exception as e:  # 모든 인코딩 실패 시
                logger.error(f"TXT 텍스트 추출 실패: {e}")
                return ""
        except Exception as e:  # 기타 예외
            logger.error(f"TXT 텍스트 추출 실패: {e}")
            return ""

    def process_docx_document(self, file_path: str, law_code: str) -> Dict:
        """DOCX 법령 문서 처리"""
        logger.info(f"DOCX 문서 처리 시작: {file_path}")
        
        # DOCX 텍스트 추출
        text_content = self._extract_docx_text(file_path)
        
        if not text_content:
            return {"success": False, "error": "Failed to extract text from DOCX file"}
        
        # 법령 정보 추출
        law_info = self._extract_law_metadata(text_content, law_code)
        
        # 조문별 분할
        articles = self._extract_articles(text_content, law_code)
        
        # Neo4j에 저장
        result = self._save_to_graph(law_info, articles)
        
        logger.info(f"DOCX 처리 완료: {len(articles)}개 조문 처리")
        return result

    def process_doc_document(self, file_path: str, law_code: str) -> Dict:
        """DOC 법령 문서 처리 (.doc 파일용)"""
        logger.info(f"DOC 문서 처리 시작: {file_path}")
        
        # DOC 텍스트 추출
        text_content = self._extract_doc_text(file_path)
        
        if not text_content:
            return {"success": False, "error": "Failed to extract text from DOC file", "articles_count": 0}
        
        # 법령 정보 추출
        law_info = self._extract_law_metadata(text_content, law_code)
        
        # 조문별 분할
        articles = self._extract_articles(text_content, law_code)
        
        # Neo4j에 저장
        result = self._save_to_graph(law_info, articles)
        
        # 성공 여부와 조문 수 추가
        if result.get("status") == "success":
            result["success"] = True
            result["articles_count"] = result.get("processed_articles", 0)
        else:
            result["success"] = False
            result["articles_count"] = 0
        
        logger.info(f"DOC 처리 완료: {len(articles)}개 조문 처리")
        return result

    def _extract_docx_text(self, file_path: str) -> str:
        """DOCX 파일에서 텍스트 추출"""
        try:
            doc = Document(file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
                
            return text_content.strip()
            
        except Exception as e:
            logger.error(f"DOCX 텍스트 추출 실패: {e}")
            return ""

    def _extract_doc_text(self, file_path: str) -> str:
        """DOC 파일에서 텍스트 추출"""
        try:
            import subprocess
            import tempfile
            import os
            import chardet
            
            # 방법 1: LibreOffice를 사용한 변환
            try:
                # 임시 디렉토리 생성
                with tempfile.TemporaryDirectory() as temp_dir:
                    # LibreOffice 명령어로 DOC를 텍스트로 변환
                    cmd = [
                        "soffice", 
                        "--headless", 
                        "--convert-to", "txt:Text", 
                        "--outdir", temp_dir,
                        file_path
                    ]
                    
                    logger.info(f"LibreOffice로 DOC 변환 시작: {file_path}")
                    result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
                    
                    if result.returncode == 0:
                        # 변환된 파일 찾기
                        base_name = os.path.splitext(os.path.basename(file_path))[0]
                        converted_file = os.path.join(temp_dir, f"{base_name}.txt")
                        
                        if os.path.exists(converted_file):
                            with open(converted_file, 'r', encoding='utf-8') as f:
                                text_content = f.read()
                            
                            if text_content and len(text_content.strip()) > 100:
                                logger.info(f"LibreOffice로 DOC 변환 성공: {len(text_content)} 문자")
                                return text_content.strip()
                    else:
                        logger.warning(f"LibreOffice DOC 변환 실패: {result.stderr}")
                        
            except Exception as e:
                logger.warning(f"LibreOffice DOC 변환 실패: {e}")
            
            # 방법 2: docx2txt 사용 (DOC/DOCX 모두 지원)
            try:
                import docx2txt
                text_content = docx2txt.process(file_path)
                
                if text_content and len(text_content.strip()) > 100:
                    logger.info(f"docx2txt로 DOC 파일 처리 성공: {len(text_content)} 문자")
                    return text_content.strip()
            except Exception as e:
                logger.warning(f"docx2txt 처리 실패: {e}")
            
            # 방법 3: python-docx로 시도 (일부 .doc 파일은 실제로 docx 형식)
            try:
                doc = Document(file_path)
                text_content = ""
                for paragraph in doc.paragraphs:
                    text_content += paragraph.text + "\n"
                
                if text_content.strip() and len(text_content.strip()) > 100:
                    logger.info("DOC 파일을 DOCX 방식으로 성공적으로 읽음")
                    return text_content.strip()
            except Exception as e:
                logger.warning(f"python-docx 처리 실패: {e}")
            
            # 방법 4: 바이너리 읽기 + 개선된 텍스트 추출 (최후 수단)
            logger.info(f"DOC 파일 {file_path}를 바이너리 모드로 읽습니다.")
            
            with open(file_path, 'rb') as file:
                content = file.read()
            
            # 인코딩 감지
            encoding_result = chardet.detect(content)
            detected_encoding = encoding_result.get('encoding', 'cp949')
            
            # 텍스트 추출 시도
            text_content = ""
            encodings_to_try = [detected_encoding, 'cp949', 'euc-kr', 'utf-8', 'utf-16', 'latin-1']
            
            for encoding in encodings_to_try:
                if not encoding:
                    continue
                try:
                    text_content = content.decode(encoding, errors='ignore')
                    if len(text_content.strip()) > 100:
                        break
                except:
                    continue
            
            if not text_content or len(text_content.strip()) < 100:
                logger.error(f"DOC 파일에서 충분한 텍스트를 추출하지 못함")
                return ""
            
            # DOC 파일의 특수 문자 및 포맷 정리
            text_content = self._clean_doc_text(text_content)
            
            logger.info(f"DOC 파일에서 {len(text_content)} 문자 추출 완료")
            return text_content.strip()
            
        except Exception as e:
            logger.error(f"DOC 텍스트 추출 실패: {e}")
            return ""
    
    def _clean_doc_text(self, text: str) -> str:
        """DOC 텍스트 정리"""
        import re
        
        # DOC 파일의 제어 문자 및 특수 포맷 정리
        # RTF 태그 제거
        text = re.sub(r'\\[a-zA-Z]+\d*\s*', ' ', text)
        text = re.sub(r'[{}]', ' ', text)
        
        # 한글, 영문, 숫자, 기본 문장부호, 법령 관련 기호만 유지
        text = re.sub(r'[^\w\s가-힣.,()①②③④⑤⑥⑦⑧⑨⑩제조항호목\-\n"\'<>【】『』「」]', ' ', text)
        
        # 연속된 공백 정리
        text = re.sub(r'\s+', ' ', text)
        
        # 연속된 줄바꿈 정리
        text = re.sub(r'\n\s*\n', '\n', text)
        
        return text

    def process_hwp_document(self, file_path: str, law_code: str) -> Dict:
        """HWP 법령 문서 처리"""
        logger.info(f"HWP 문서 처리 시작: {file_path}")
        
        # HWP 텍스트 추출
        text_content = self._extract_hwp_text(file_path)
        
        if not text_content:
            return {"success": False, "error": "Failed to extract text from HWP file"}
        
        # 법령 정보 추출
        law_info = self._extract_law_metadata(text_content, law_code)
        
        # 조문별 분할
        articles = self._extract_articles(text_content, law_code)
        
        # Neo4j에 저장
        result = self._save_to_graph(law_info, articles)
        
        logger.info(f"HWP 처리 완료: {len(articles)}개 조문 처리")
        return result

    def _extract_hwp_text(self, file_path: str) -> str:
        """HWP 파일에서 텍스트 추출 - 다중 방법 시도"""
        logger.info(f"HWP 파일 처리 시작: {file_path}")
        
        # 방법 1: LibreOffice를 통한 변환 (최우선)
        try:
            text_content = self._extract_hwp_with_libreoffice(file_path)
            if text_content and len(text_content.strip()) > 100:
                logger.info(f"LibreOffice로 HWP 처리 성공: {len(text_content)} 문자")
                return text_content.strip()
        except Exception as e:
            logger.warning(f"LibreOffice 처리 실패: {e}")
        
        # 방법 2: olefile을 이용한 직접 파싱
        try:
            text_content = self._extract_hwp_with_olefile(file_path)
            if text_content and len(text_content.strip()) > 100:
                logger.info(f"olefile로 HWP 처리 성공: {len(text_content)} 문자")
                return text_content.strip()
        except Exception as e:
            logger.warning(f"olefile 처리 실패: {e}")
        
        # 방법 3: 개선된 바이너리 방식
        try:
            text_content = self._extract_hwp_binary_advanced(file_path)
            if text_content and len(text_content.strip()) > 50:
                logger.info(f"개선된 바이너리 방식으로 HWP 처리 완료: {len(text_content)} 문자")
                return text_content.strip()
        except Exception as e:
            logger.warning(f"개선된 바이너리 처리 실패: {e}")
        
        # 방법 4: 기본 바이너리 방식 (백업)
        try:
            text_content = self._extract_hwp_binary_fallback(file_path)
            if text_content and len(text_content.strip()) > 50:
                logger.info(f"기본 바이너리 방식으로 HWP 처리 완료: {len(text_content)} 문자")
                return text_content.strip()
        except Exception as e:
            logger.warning(f"기본 바이너리 처리 실패: {e}")
        
        logger.error(f"모든 HWP 처리 방법 실패: {file_path}")
        return ""
    
    def _extract_hwp_binary_advanced(self, file_path: str) -> str:
        """개선된 HWP 바이너리 처리"""
        try:
            import chardet
            import struct
            
            with open(file_path, 'rb') as file:
                content = file.read()
            
            # HWP 파일 시그니처 확인
            if content[:4] != b'\xd0\xcf\x11\xe0':
                logger.warning("HWP 파일 시그니처가 올바르지 않음")
            
            # 텍스트 패턴 검색
            text_content = ""
            
            # 한글 유니코드 범위에서 텍스트 추출
            korean_pattern = re.compile(r'[\uac00-\ud7af]+')
            
            # 여러 인코딩으로 디코딩 시도
            encodings = ['utf-16le', 'utf-16be', 'cp949', 'euc-kr', 'utf-8']
            
            for encoding in encodings:
                try:
                    decoded = content.decode(encoding, errors='ignore')
                    korean_matches = korean_pattern.findall(decoded)
                    
                    if korean_matches and len(''.join(korean_matches)) > 100:
                        # 한글 텍스트가 충분히 발견되면
                        text_content = decoded
                        logger.info(f"인코딩 {encoding}으로 한글 텍스트 발견")
                        break
                except:
                    continue
            
            if not text_content:
                # 바이트 패턴으로 한글 텍스트 직접 추출
                text_content = self._extract_korean_from_bytes(content)
            
            if text_content:
                # 텍스트 정리
                text_content = self._clean_hwp_text(text_content)
                
                # 법령 관련 키워드가 있는지 확인
                legal_keywords = ['조', '항', '호', '법', '령', '규칙', '조례', '시행']
                if any(keyword in text_content for keyword in legal_keywords):
                    return text_content
            
            return ""
            
        except Exception as e:
            logger.error(f"개선된 HWP 바이너리 처리 실패: {e}")
            return ""
    
    def _extract_korean_from_bytes(self, content: bytes) -> str:
        """바이트에서 직접 한글 텍스트 추출"""
        try:
            text_parts = []
            
            # UTF-16LE 패턴으로 한글 검색
            i = 0
            while i < len(content) - 1:
                try:
                    # 2바이트씩 읽어서 한글 유니코드 범위 확인
                    char_code = struct.unpack('<H', content[i:i+2])[0]
                    
                    # 한글 유니코드 범위 (가-힣: 0xAC00-0xD7AF)
                    if 0xAC00 <= char_code <= 0xD7AF:
                        # 연속된 한글 문자들 추출
                        korean_text = ""
                        j = i
                        while j < len(content) - 1:
                            char_code = struct.unpack('<H', content[j:j+2])[0]
                            if 0xAC00 <= char_code <= 0xD7AF:
                                korean_text += chr(char_code)
                                j += 2
                            elif char_code in [0x0020, 0x000A, 0x000D]:  # 공백, 줄바꿈
                                korean_text += chr(char_code)
                                j += 2
                            else:
                                break
                        
                        if len(korean_text) > 2:
                            text_parts.append(korean_text)
                        i = j
                    else:
                        i += 1
                except:
                    i += 1
            
            return ' '.join(text_parts)
            
        except Exception as e:
            logger.warning(f"바이트에서 한글 추출 실패: {e}")
            return ""

    def _extract_hwp_with_libreoffice(self, file_path: str) -> str:
        """LibreOffice를 사용한 개선된 HWP 변환"""
        try:
            import subprocess
            import tempfile
            import os
            
            # 임시 디렉토리 생성
            with tempfile.TemporaryDirectory() as temp_dir:
                # 여러 변환 형식 시도
                formats_to_try = [
                    ("txt:Text", ".txt"),
                    ("odt", ".odt"),
                    ("docx", ".docx"),
                    ("rtf", ".rtf")
                ]
                
                for format_name, extension in formats_to_try:
                    try:
                        # LibreOffice 명령어로 변환
                        cmd = [
                            "soffice", 
                            "--headless", 
                            "--convert-to", format_name,
                            "--outdir", temp_dir,
                            file_path
                        ]
                        
                        logger.info(f"LibreOffice {format_name} 변환 시도: {file_path}")
                        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
                        
                        if result.returncode == 0:
                            # 변환된 파일 찾기
                            base_name = os.path.splitext(os.path.basename(file_path))[0]
                            converted_file = os.path.join(temp_dir, f"{base_name}{extension}")
                            
                            if os.path.exists(converted_file):
                                # 변환된 파일에서 텍스트 추출
                                if extension == ".txt":
                                    with open(converted_file, 'r', encoding='utf-8') as f:
                                        text_content = f.read()
                                elif extension == ".docx":
                                    text_content = self._extract_docx_text(converted_file)
                                elif extension == ".odt":
                                    text_content = self._extract_odt_text(converted_file)
                                elif extension == ".rtf":
                                    text_content = self._extract_rtf_text(converted_file)
                                else:
                                    continue
                                
                                if text_content and len(text_content.strip()) > 100:
                                    logger.info(f"LibreOffice {format_name} 변환 성공")
                                    return self._clean_hwp_text(text_content)
                    except Exception as e:
                        logger.warning(f"LibreOffice {format_name} 변환 실패: {e}")
                        continue
                
                return ""
                
        except Exception as e:
            logger.error(f"LibreOffice 변환 실패: {e}")
            return ""
    
    def _extract_hwp_with_olefile(self, file_path: str) -> str:
        """olefile을 사용한 HWP 직접 파싱"""
        try:
            import olefile
            
            # HWP 파일이 OLE 형식인지 확인
            if not olefile.isOleFile(file_path):
                logger.warning("HWP 파일이 OLE 형식이 아님")
                return ""
            
            # OLE 파일 열기
            with olefile.OleFileIO(file_path) as ole:
                text_content = ""
                
                # HWP 5.0 형식의 스트림 탐색
                streams_to_check = [
                    'BodyText/Section0',
                    'BodyText/Section1',
                    'BodyText/Section2',
                    'BodyText/Section3',
                    'BodyText/Section4',
                    'BodyText/Section5',
                    'DocOptions',
                    'BinData',
                    'PreviewText',
                    'Scripts',
                    'DocInfo'
                ]
                
                # 실제 존재하는 스트림 목록 확인
                available_streams = ole.listdir()
                logger.info(f"HWP 파일 내 스트림: {available_streams}")
                
                # 각 스트림에서 텍스트 추출 시도
                for stream_path in streams_to_check:
                    try:
                        if ole._olestream_size(stream_path) > 0:
                            # 스트림 데이터 읽기
                            stream_data = ole._olestream_read(stream_path)
                            
                            # 텍스트 추출 시도 (여러 인코딩)
                            for encoding in ['utf-8', 'cp949', 'euc-kr', 'utf-16']:
                                try:
                                    decoded_text = stream_data.decode(encoding, errors='ignore')
                                    
                                    # 한글 문자가 포함된 의미있는 텍스트인지 확인
                                    if any('\uac00' <= char <= '\ud7af' for char in decoded_text[:1000]):
                                        # 제어 문자 제거 및 정리
                                        clean_text = self._clean_hwp_stream_text(decoded_text)
                                        if len(clean_text.strip()) > 50:
                                            text_content += clean_text + "\n"
                                            break
                                except:
                                    continue
                    except:
                        continue
                
                # BodyText 섹션들 특별 처리
                for i in range(10):  # Section0 ~ Section9 확인
                    section_name = f'BodyText/Section{i}'
                    try:
                        if ole._olestream_size(section_name) > 0:
                            section_data = ole._olestream_read(section_name)
                            
                            # HWP 바이너리 구조 파싱 시도
                            parsed_text = self._parse_hwp_bodytext_section(section_data)
                            if parsed_text:
                                text_content += parsed_text + "\n"
                    except:
                        continue
                
                return self._clean_hwp_text(text_content) if text_content else ""
                
        except ImportError:
            logger.warning("olefile 라이브러리가 설치되지 않음")
            return ""
        except Exception as e:
            logger.error(f"olefile HWP 파싱 실패: {e}")
            return ""
    
    def _parse_hwp_bodytext_section(self, section_data: bytes) -> str:
        """HWP BodyText 섹션 바이너리 데이터 파싱"""
        try:
            text_content = ""
            
            # HWP 5.0 레코드 구조 파싱
            offset = 0
            while offset < len(section_data) - 8:
                try:
                    # 레코드 헤더 읽기 (8바이트)
                    if offset + 8 > len(section_data):
                        break
                    
                    record_header = section_data[offset:offset+8]
                    record_type = int.from_bytes(record_header[:4], byteorder='little')
                    record_size = int.from_bytes(record_header[4:8], byteorder='little')
                    
                    # 레코드 크기 검증
                    if record_size <= 0 or record_size > len(section_data) - offset:
                        offset += 1
                        continue
                    
                    # 레코드 데이터 읽기
                    record_data = section_data[offset+8:offset+8+record_size]
                    
                    # 텍스트 레코드 타입 확인 (일반적으로 67, 68 등)
                    if record_type in [67, 68, 69, 70]:  # 텍스트 관련 레코드
                        # UTF-16LE로 디코딩 시도
                        try:
                            text = record_data.decode('utf-16le', errors='ignore')
                            # null 문자 제거 및 정리
                            text = text.replace('\x00', '').strip()
                            if len(text) > 2 and any('\uac00' <= char <= '\ud7af' for char in text):
                                text_content += text + " "
                        except:
                            pass
                    
                    offset += 8 + record_size
                    
                except Exception:
                    offset += 1  # 다음 바이트로 이동
                    continue
            
            return text_content.strip()
            
        except Exception as e:
            logger.warning(f"HWP 섹션 파싱 실패: {e}")
            return ""
    
    def _clean_hwp_stream_text(self, text: str) -> str:
        """HWP 스트림에서 추출한 텍스트 정리"""
        import re
        
        # 제어 문자 및 특수 기호 제거
        text = re.sub(r'[\x00-\x1f\x7f-\x9f]', ' ', text)
        
        # HWP 관련 태그 및 메타데이터 제거
        text = re.sub(r'<[^>]+>', ' ', text)
        text = re.sub(r'\\[a-zA-Z]+\d*\s*', ' ', text)
        
        # 연속된 공백 정리
        text = re.sub(r'\s+', ' ', text)
        
        # 의미있는 한글 텍스트만 추출
        lines = text.split('\n')
        clean_lines = []
        
        for line in lines:
            line = line.strip()
            # 한글이 포함되고 최소 길이를 만족하는 라인만 유지
            if len(line) > 5 and any('\uac00' <= char <= '\ud7af' for char in line):
                clean_lines.append(line)
        
        return '\n'.join(clean_lines)
    
    def _extract_hwp_binary_fallback(self, file_path: str) -> str:
        """기존 바이너리 방식 (개선된 버전)"""
        try:
            import chardet
            
            with open(file_path, 'rb') as file:
                content = file.read()
            
            # 인코딩 감지
            encoding_result = chardet.detect(content)
            detected_encoding = encoding_result.get('encoding', 'cp949')
            
            # 텍스트 추출 시도
            text_content = ""
            encodings_to_try = [detected_encoding, 'cp949', 'euc-kr', 'utf-8', 'utf-16', 'utf-16le']
            
            for encoding in encodings_to_try:
                if not encoding:
                    continue
                try:
                    text_content = content.decode(encoding, errors='ignore')
                    # 한글 텍스트가 포함되어 있는지 확인
                    if any('\uac00' <= char <= '\ud7af' for char in text_content[:1000]):
                        break
                except:
                    continue
            
            if not text_content:
                logger.error("HWP 파일에서 텍스트를 추출할 수 없습니다.")
                return ""
            
            # HWP 텍스트 정리
            text_content = self._clean_hwp_text(text_content)
            
            return text_content
            
        except Exception as e:
            logger.error(f"HWP 바이너리 추출 실패: {e}")
            return ""
    
    def _extract_odt_text(self, file_path: str) -> str:
        """ODT 파일에서 텍스트 추출"""
        try:
            import zipfile
            from xml.etree import ElementTree as ET
            
            text_content = ""
            
            with zipfile.ZipFile(file_path, 'r') as odt_zip:
                # content.xml 파일에서 텍스트 추출
                if 'content.xml' in odt_zip.namelist():
                    content_xml = odt_zip.read('content.xml')
                    root = ET.fromstring(content_xml)
                    
                    # 모든 텍스트 노드 추출
                    for elem in root.iter():
                        if elem.text:
                            text_content += elem.text + " "
                        if elem.tail:
                            text_content += elem.tail + " "
            
            return text_content.strip()
            
        except Exception as e:
            logger.warning(f"ODT 텍스트 추출 실패: {e}")
            return ""
    
    def _extract_rtf_text(self, file_path: str) -> str:
        """RTF 파일에서 텍스트 추출"""
        try:
            import re
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                rtf_content = file.read()
            
            # RTF 태그 제거
            text_content = re.sub(r'\\[a-zA-Z]+\d*\s*', ' ', rtf_content)
            text_content = re.sub(r'[{}]', ' ', text_content)
            text_content = re.sub(r'\s+', ' ', text_content)
            
            return text_content.strip()
            
        except Exception as e:
            logger.warning(f"RTF 텍스트 추출 실패: {e}")
            return ""

    def _clean_hwp_text(self, text: str) -> str:
        """HWP 텍스트 정리"""
        import re
        
        # DOC 파일의 제어 문자 및 특수 포맷 정리
        # RTF 태그 제거
        text = re.sub(r'\\[a-zA-Z]+\d*\s*', ' ', text)
        text = re.sub(r'[{}]', ' ', text)
        
        # 한글, 영문, 숫자, 기본 문장부호, 법령 관련 기호만 유지
        text = re.sub(r'[^\w\s가-힣.,()①②③④⑤⑥⑦⑧⑨⑩제조항호목\-\n"\'<>【】『』「」]', ' ', text)
        
        # 연속된 공백 정리
        text = re.sub(r'\s+', ' ', text)
        
        # 연속된 줄바꿈 정리
        text = re.sub(r'\n\s*\n', '\n', text)
        
        return text


# 사용 예시
if __name__ == "__main__":
    # 그래프 관리자 초기화
    graph_manager = LegalGraphManager()
    
    # 문서 처리기 초기화
    processor = LegalDocumentProcessor(graph_manager)
    
    # 샘플 문서 처리
    # result = processor.process_pdf_document("data/laws/도시정비법_전문.pdf", "URBAN_REDEVELOPMENT_LAW")
    # print(f"처리 결과: {result}")
    
    # 일괄 처리
    results = processor.process_documents_batch("data")
    print(f"일괄 처리 결과: {results}")
    
    # 연결 종료
    graph_manager.close() 